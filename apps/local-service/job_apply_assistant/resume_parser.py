from __future__ import annotations

import re
from pathlib import Path
from uuid import uuid4

from docx import Document
from pypdf import PdfReader

from job_apply_assistant.models import ResumeProfile


SKILL_KEYWORDS = [
    "Python",
    "TypeScript",
    "JavaScript",
    "React",
    "ROS",
    "机器人",
    "机械臂",
    "算法",
    "控制",
    "感知",
    "深度学习",
    "机器学习",
]


class ResumeParser:
    def parse(self, path: Path) -> ResumeProfile:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text = self._extract_pdf(path)
        elif suffix == ".docx":
            text = self._extract_docx(path)
        else:
            file_type = suffix or "no extension"
            raise ValueError(f"Unsupported resume file type: {file_type}")

        cleaned = self._normalize_text(text)
        skills = [skill for skill in SKILL_KEYWORDS if skill.lower() in cleaned.lower()]

        return ResumeProfile(
            id=f"resume_{uuid4().hex}",
            fileName=path.name,
            rawText=cleaned,
            summary=cleaned[:500],
            skills=skills,
            yearsOfExperience=self._extract_years(cleaned),
            projectHighlights=self._extract_project_highlights(cleaned),
            education=self._extract_education(cleaned),
            targetRoleSuggestions=self._suggest_roles(cleaned),
        )

    def _extract_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _extract_docx(self, path: Path) -> str:
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _extract_years(self, text: str) -> float:
        match = re.search(r"(\d+(?:\.\d+)?)\s*(?:年|years?|yrs?)\b", text, flags=re.IGNORECASE)
        return float(match.group(1)) if match else 0.0

    def _extract_project_highlights(self, text: str) -> list[str]:
        sentences = re.split(r"[。.!?]", text)
        return [
            sentence.strip()
            for sentence in sentences
            if "项目" in sentence or "project" in sentence.lower()
        ][:5]

    def _extract_education(self, text: str) -> list[str]:
        education_words = ["本科", "硕士", "博士", "大专", "Bachelor", "Master", "PhD"]
        return [word for word in education_words if word.lower() in text.lower()]

    def _suggest_roles(self, text: str) -> list[str]:
        roles: list[str] = []
        if "机器人" in text or "ROS" in text:
            roles.append("机器人软件工程师")
        if "算法" in text or "深度学习" in text:
            roles.append("算法工程师")
        if "React" in text or "TypeScript" in text:
            roles.append("前端工程师")
        return roles or ["软件工程师"]
