from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas

from job_apply_assistant.resume_parser import ResumeParser


def write_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def write_docx_table(path: Path, rows: list[tuple[str, str]]) -> None:
    document = Document()
    table = document.add_table(rows=len(rows), cols=2)
    for row_index, row in enumerate(rows):
        for cell_index, value in enumerate(row):
            table.cell(row_index, cell_index).text = value
    document.save(path)


def write_pdf(path: Path, text: str) -> None:
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 720, text)
    pdf.save()


def test_parse_docx_resume(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    write_docx(path, "张三 \n 机器人算法工程师   Python ROS 机械臂 3年 本科 项目经验")

    profile = ResumeParser().parse(path)

    assert profile.file_name == "resume.docx"
    assert profile.raw_text == "张三 机器人算法工程师 Python ROS 机械臂 3年 本科 项目经验"
    assert "机器人算法工程师" in profile.raw_text
    assert "Python" in profile.skills
    assert profile.years_of_experience == 3.0
    assert profile.project_highlights == ["张三 机器人算法工程师 Python ROS 机械臂 3年 本科 项目经验"]
    assert profile.education == ["本科"]
    assert profile.target_role_suggestions == ["机器人软件工程师", "算法工程师"]

    payload = profile.to_wire()
    assert payload["fileName"] == "resume.docx"
    assert payload["rawText"] == profile.raw_text
    assert payload["yearsOfExperience"] == 3.0
    assert payload["targetRoleSuggestions"] == ["机器人软件工程师", "算法工程师"]


def test_parse_docx_resume_table_text(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    write_docx_table(
        path,
        [
            ("Name", "Li Engineer"),
            ("Experience", "5 years experience in Python Robot Control project"),
            ("Education", "Bachelor of Robotics"),
        ],
    )

    profile = ResumeParser().parse(path)

    assert "Li Engineer" in profile.raw_text
    assert "Python Robot Control project" in profile.raw_text
    assert profile.years_of_experience == 5.0
    assert len(profile.project_highlights) == 1
    assert "5 years experience in Python Robot Control project" in profile.project_highlights[0]
    assert profile.education == ["Bachelor"]


def test_parse_chinese_years_followed_by_experience_text(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    write_docx(path, "拥有3年机器人项目经验 Python ROS")

    profile = ResumeParser().parse(path)

    assert profile.years_of_experience == 3.0


def test_parse_chinese_decimal_years_followed_by_project_text(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    write_docx(path, "拥有3.5年项目经验 Python ROS")

    profile = ResumeParser().parse(path)

    assert profile.years_of_experience == 3.5


def test_parse_pdf_resume(tmp_path: Path) -> None:
    path = tmp_path / "resume.pdf"
    write_pdf(path, "Li Engineer Python Robot Control project Master 4 yrs")

    profile = ResumeParser().parse(path)

    assert profile.file_name == "resume.pdf"
    assert "Python" in profile.raw_text
    assert profile.years_of_experience == 4.0
    assert profile.project_highlights == ["Li Engineer Python Robot Control project Master 4 yrs"]
    assert profile.education == ["Master"]


def test_reject_unsupported_file(tmp_path: Path) -> None:
    path = tmp_path / "resume.txt"
    path.write_text("plain text", encoding="utf-8")

    try:
        ResumeParser().parse(path)
    except ValueError as exc:
        assert "Unsupported resume file type" in str(exc)
    else:
        raise AssertionError("unsupported file should raise ValueError")


def test_reject_empty_docx_resume(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    write_docx(path, "")

    try:
        ResumeParser().parse(path)
    except ValueError as exc:
        assert "No extractable resume text found" in str(exc)
    else:
        raise AssertionError("empty resume should raise ValueError")


def test_reject_corrupt_pdf_with_parser_error(tmp_path: Path) -> None:
    path = tmp_path / "resume.pdf"
    path.write_bytes(b"not a pdf")

    try:
        ResumeParser().parse(path)
    except ValueError as exc:
        message = str(exc)
        assert "resume.pdf" in message
        assert "Could not extract resume text" in message
    else:
        raise AssertionError("corrupt PDF should raise ValueError")


def test_reject_corrupt_docx_with_parser_error(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    path.write_bytes(b"not a docx")

    try:
        ResumeParser().parse(path)
    except ValueError as exc:
        message = str(exc)
        assert "resume.docx" in message
        assert "Could not extract resume text" in message
    else:
        raise AssertionError("corrupt DOCX should raise ValueError")


def test_reject_file_without_extension(tmp_path: Path) -> None:
    path = tmp_path / "resume"
    path.write_text("plain text", encoding="utf-8")

    try:
        ResumeParser().parse(path)
    except ValueError as exc:
        assert "Unsupported resume file type: no extension" in str(exc)
    else:
        raise AssertionError("unsupported file should raise ValueError")
